"""
Word Document Operations Service using python-docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from pathlib import Path
from typing import List, Optional, Dict
import os

from app.config import settings
from app.utils.helpers import generate_file_id


class WordService:
    """Service for Word document manipulation operations"""
    
    @staticmethod
    def get_document_info(file_path: Path) -> dict:
        """Get Word document information"""
        doc = Document(file_path)
        
        # Count paragraphs and tables
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        section_count = len(doc.sections)
        
        # Get core properties
        props = doc.core_properties
        
        return {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "section_count": section_count,
            "author": props.author,
            "title": props.title,
            "subject": props.subject,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "file_size": file_path.stat().st_size
        }
    
    @staticmethod
    def merge_documents(file_paths: List[Path], output_filename: str = "merged.docx") -> Path:
        """
        Merge multiple Word documents into one
        """
        if not file_paths:
            raise ValueError("No files provided for merging")
        
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_{output_filename}"
        
        # Start with the first document
        merged_doc = Document(file_paths[0])
        
        # Append content from other documents
        for file_path in file_paths[1:]:
            # Add page break before new document content
            merged_doc.add_page_break()
            
            doc = Document(file_path)
            
            # Copy paragraphs
            for para in doc.paragraphs:
                new_para = merged_doc.add_paragraph()
                new_para.style = para.style
                new_para.alignment = para.alignment
                
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.name:
                        new_run.font.name = run.font.name
            
            # Copy tables
            for table in doc.tables:
                # Create new table with same dimensions
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = table.style
                
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
        
        merged_doc.save(output_path)
        return output_path
    
    @staticmethod
    def split_by_sections(file_path: Path) -> List[Path]:
        """
        Split document by sections
        """
        doc = Document(file_path)
        output_paths = []
        
        # For each section, create a new document
        for i, section in enumerate(doc.sections):
            new_doc = Document()
            
            # Copy section properties
            new_section = new_doc.sections[0]
            new_section.page_width = section.page_width
            new_section.page_height = section.page_height
            new_section.orientation = section.orientation
            
            output_filename = f"{generate_file_id()}_section_{i+1}.docx"
            output_path = settings.OUTPUT_DIR / output_filename
            
            new_doc.save(output_path)
            output_paths.append(output_path)
        
        return output_paths
    
    @staticmethod
    def extract_text(file_path: Path) -> str:
        """
        Extract all text from document
        """
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        return '\n'.join(full_text)
    
    @staticmethod
    def add_header(file_path: Path, text: str, align: str = "center") -> Path:
        """
        Add header to all sections
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_with_header.docx"
        
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT
        }
        
        for section in doc.sections:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.text = text
            paragraph.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def add_footer(file_path: Path, text: str, include_page_number: bool = False) -> Path:
        """
        Add footer to all sections
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_with_footer.docx"
        
        for section in doc.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.text = text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def add_watermark(file_path: Path, text: str) -> Path:
        """
        Add text watermark to document
        Note: python-docx has limited watermark support
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_watermarked.docx"
        
        # Add watermark text to header (simplified approach)
        for section in doc.sections:
            header = section.header
            paragraph = header.add_paragraph()
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.size = Pt(48)
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def insert_image(file_path: Path, image_path: Path, 
                     width_inches: float = 4.0, position: int = 0) -> Path:
        """
        Insert image into document
        position: paragraph index where to insert (0 = beginning)
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_with_image.docx"
        
        # Insert image at specified position
        paragraph = doc.paragraphs[position] if position < len(doc.paragraphs) else doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def insert_table(file_path: Path, data: List[List[str]], 
                     headers: Optional[List[str]] = None, position: int = -1) -> Path:
        """
        Insert table into document
        data: 2D list of cell values
        headers: Optional header row
        position: paragraph index where to insert (-1 = end)
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_with_table.docx"
        
        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else (len(headers) if headers else 1)
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Add headers if provided
        start_row = 0
        if headers:
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            start_row = 1
        
        # Add data
        for i, row_data in enumerate(data):
            for j, cell_value in enumerate(row_data):
                table.rows[i + start_row].cells[j].text = str(cell_value)
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def replace_text(file_path: Path, find_text: str, replace_text: str) -> Path:
        """
        Find and replace text in document
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_replaced.docx"
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if find_text in para.text:
                for run in para.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if find_text in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if find_text in run.text:
                                    run.text = run.text.replace(find_text, replace_text)
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def update_styles(file_path: Path, style_updates: Dict) -> Path:
        """
        Update document styles
        style_updates: dict with font_name, font_size, line_spacing, etc.
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_styled.docx"
        
        font_name = style_updates.get('font_name')
        font_size = style_updates.get('font_size')
        
        for para in doc.paragraphs:
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def set_page_orientation(file_path: Path, orientation: str = "portrait") -> Path:
        """
        Set page orientation for all sections
        orientation: "portrait" or "landscape"
        """
        doc = Document(file_path)
        output_path = settings.OUTPUT_DIR / f"{generate_file_id()}_oriented.docx"
        
        for section in doc.sections:
            if orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap width and height
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height
            else:
                section.orientation = WD_ORIENT.PORTRAIT
        
        doc.save(output_path)
        return output_path


# Create singleton instance
word_service = WordService()
