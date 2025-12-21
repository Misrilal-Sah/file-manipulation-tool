"""
Comprehensive test script for Document Manipulation Tool
Tests all PDF, Word, and conversion features
"""
import requests
import os
import json
from pathlib import Path

BASE_URL = "http://127.0.0.1:8000"
TEST_DIR = Path(__file__).parent / "test_files"
TEST_DIR.mkdir(exist_ok=True)

# Results tracking
results = {"passed": [], "failed": [], "skipped": []}


def log_result(test_name: str, passed: bool, message: str = ""):
    status = "✅ PASS" if passed else "❌ FAIL"
    print(f"{status}: {test_name}")
    if message:
        print(f"   └── {message}")
    if passed:
        results["passed"].append(test_name)
    else:
        results["failed"].append((test_name, message))


def create_test_pdf():
    """Create a simple test PDF using PyMuPDF"""
    try:
        import fitz
        pdf_path = TEST_DIR / "test.pdf"
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((50, 50), f"Test Page {i+1}", fontsize=24)
            page.insert_text((50, 100), "This is a test document for API testing.", fontsize=12)
        doc.save(str(pdf_path))
        doc.close()
        return pdf_path
    except Exception as e:
        print(f"Error creating test PDF: {e}")
        return None


def create_test_word():
    """Create a simple test Word document"""
    try:
        from docx import Document
        doc_path = TEST_DIR / "test.docx"
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph for API testing.")
        doc.add_paragraph("Second paragraph with more content.")
        doc.save(str(doc_path))
        return doc_path
    except Exception as e:
        print(f"Error creating test Word doc: {e}")
        return None


def test_health():
    """Test health endpoint"""
    try:
        r = requests.get(f"{BASE_URL}/health")
        log_result("Health Check", r.status_code == 200, f"Status: {r.json()}")
        return r.status_code == 200
    except Exception as e:
        log_result("Health Check", False, str(e))
        return False


def test_upload_pdf(pdf_path: Path):
    """Test PDF upload"""
    try:
        with open(pdf_path, 'rb') as f:
            r = requests.post(
                f"{BASE_URL}/api/upload/",
                files={"file": ("test.pdf", f, "application/pdf")}
            )
        if r.status_code == 200:
            data = r.json()
            log_result("Upload PDF", True, f"File ID: {data.get('file_id')}")
            return data.get('file_id')
        else:
            log_result("Upload PDF", False, r.text)
            return None
    except Exception as e:
        log_result("Upload PDF", False, str(e))
        return None


def test_upload_word(doc_path: Path):
    """Test Word document upload"""
    try:
        with open(doc_path, 'rb') as f:
            r = requests.post(
                f"{BASE_URL}/api/upload/",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        if r.status_code == 200:
            data = r.json()
            log_result("Upload Word", True, f"File ID: {data.get('file_id')}")
            return data.get('file_id')
        else:
            log_result("Upload Word", False, r.text)
            return None
    except Exception as e:
        log_result("Upload Word", False, str(e))
        return None


def test_pdf_info(file_id: str):
    """Test PDF info endpoint"""
    try:
        r = requests.get(f"{BASE_URL}/api/pdf/info/{file_id}")
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Info", True, f"Pages: {data.get('page_count')}")
            return True
        else:
            log_result("PDF Info", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Info", False, str(e))
        return False


def test_pdf_thumbnails(file_id: str):
    """Test PDF thumbnails endpoint"""
    try:
        r = requests.get(f"{BASE_URL}/api/pdf/thumbnails/{file_id}")
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Thumbnails", True, f"Generated {len(data.get('thumbnails', []))} thumbnails")
            return True
        else:
            log_result("PDF Thumbnails", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Thumbnails", False, str(e))
        return False


def test_pdf_merge(file_id: str):
    """Test PDF merge - needs 2 PDFs, so upload same file twice"""
    try:
        # Upload another PDF
        pdf_path = TEST_DIR / "test.pdf"
        with open(pdf_path, 'rb') as f:
            r = requests.post(
                f"{BASE_URL}/api/upload/",
                files={"file": ("test2.pdf", f, "application/pdf")}
            )
        file_id2 = r.json().get('file_id')
        
        # Merge
        r = requests.post(
            f"{BASE_URL}/api/pdf/merge",
            json={"file_ids": [file_id, file_id2], "output_filename": "merged.pdf"}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Merge", True, f"Output: {data.get('output_filename')}")
            return data.get('output_file_id')
        else:
            log_result("PDF Merge", False, r.text)
            return None
    except Exception as e:
        log_result("PDF Merge", False, str(e))
        return None


def test_pdf_split(file_id: str):
    """Test PDF split"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/split",
            json={"file_id": file_id, "ranges": [{"start": 1, "end": 1}, {"start": 2, "end": 3}]}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Split", True, f"Created {len(data.get('outputs', []))} files")
            return True
        else:
            log_result("PDF Split", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Split", False, str(e))
        return False


def test_pdf_rotate(file_id: str):
    """Test PDF rotate"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/rotate",
            json={"file_id": file_id, "pages": [1], "angle": 90}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Rotate", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("PDF Rotate", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Rotate", False, str(e))
        return False


def test_pdf_compress(file_id: str):
    """Test PDF compress"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/compress",
            json={"file_id": file_id, "quality": 80}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Compress", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("PDF Compress", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Compress", False, str(e))
        return False


def test_pdf_watermark(file_id: str):
    """Test PDF watermark"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/watermark",
            json={"file_id": file_id, "text": "CONFIDENTIAL", "opacity": 0.3, "position": "center"}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Watermark", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("PDF Watermark", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Watermark", False, str(e))
        return False


def test_pdf_protect(file_id: str):
    """Test PDF password protection"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/protect",
            json={"file_id": file_id, "password": "test123"}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Protect", True, f"Output: {data.get('output_filename')}")
            return data.get('output_file_id')
        else:
            log_result("PDF Protect", False, r.text)
            return None
    except Exception as e:
        log_result("PDF Protect", False, str(e))
        return None


def test_pdf_unlock(file_id: str):
    """Test PDF unlock"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/unlock",
            json={"file_id": file_id, "password": "test123"}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Unlock", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("PDF Unlock", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Unlock", False, str(e))
        return False


def test_pdf_extract(file_id: str):
    """Test PDF extract pages"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/extract",
            json={"file_id": file_id, "pages": [1, 2]}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Extract", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("PDF Extract", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Extract", False, str(e))
        return False


def test_pdf_delete(file_id: str):
    """Test PDF delete pages"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/pdf/delete",
            json={"file_id": file_id, "pages": [1]}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF Delete", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("PDF Delete", False, r.text)
            return False
    except Exception as e:
        log_result("PDF Delete", False, str(e))
        return False


def test_word_info(file_id: str):
    """Test Word info endpoint"""
    try:
        r = requests.get(f"{BASE_URL}/api/word/info/{file_id}")
        if r.status_code == 200:
            data = r.json()
            log_result("Word Info", True, f"Paragraphs: {data.get('paragraph_count')}")
            return True
        else:
            log_result("Word Info", False, r.text)
            return False
    except Exception as e:
        log_result("Word Info", False, str(e))
        return False


def test_word_header(file_id: str):
    """Test Word add header"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/word/add-header",
            json={"file_id": file_id, "text": "Document Header", "align": "center"}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("Word Header", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("Word Header", False, r.text)
            return False
    except Exception as e:
        log_result("Word Header", False, str(e))
        return False


def test_word_footer(file_id: str):
    """Test Word add footer"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/word/add-footer",
            json={"file_id": file_id, "text": "Page Footer", "include_page_number": True}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("Word Footer", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("Word Footer", False, r.text)
            return False
    except Exception as e:
        log_result("Word Footer", False, str(e))
        return False


def test_word_replace(file_id: str):
    """Test Word find and replace"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/word/replace-text",
            json={"file_id": file_id, "find_text": "test", "replace_text": "REPLACED"}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("Word Replace", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("Word Replace", False, r.text)
            return False
    except Exception as e:
        log_result("Word Replace", False, str(e))
        return False


def test_convert_capabilities():
    """Test conversion capabilities endpoint"""
    try:
        r = requests.get(f"{BASE_URL}/api/convert/capabilities")
        if r.status_code == 200:
            data = r.json()
            log_result("Conversion Capabilities", True, f"LibreOffice: {data.get('libreoffice_available')}")
            return data
        else:
            log_result("Conversion Capabilities", False, r.text)
            return None
    except Exception as e:
        log_result("Conversion Capabilities", False, str(e))
        return None


def test_word_to_pdf(file_id: str):
    """Test Word to PDF conversion"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/convert/word-to-pdf",
            json={"file_id": file_id}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("Word to PDF", True, f"Output: {data.get('output_filename')}")
            return True
        else:
            log_result("Word to PDF", False, r.text)
            return False
    except Exception as e:
        log_result("Word to PDF", False, str(e))
        return False


def test_pdf_to_images(file_id: str):
    """Test PDF to images conversion"""
    try:
        r = requests.post(
            f"{BASE_URL}/api/convert/pdf-to-images",
            json={"file_id": file_id, "dpi": 150, "image_format": "png"}
        )
        if r.status_code == 200:
            data = r.json()
            log_result("PDF to Images", True, f"Created {len(data.get('outputs', []))} images")
            return True
        else:
            log_result("PDF to Images", False, r.text)
            return False
    except Exception as e:
        log_result("PDF to Images", False, str(e))
        return False


def test_ocr_status():
    """Test OCR status endpoint"""
    try:
        r = requests.get(f"{BASE_URL}/api/convert/ocr/status")
        if r.status_code == 200:
            data = r.json()
            log_result("OCR Status", True, f"Available: {data.get('available')}")
            return data.get('available')
        else:
            log_result("OCR Status", False, r.text)
            return False
    except Exception as e:
        log_result("OCR Status", False, str(e))
        return False


def main():
    print("=" * 60)
    print("Document Manipulation Tool - Comprehensive Test Suite")
    print("=" * 60)
    print()
    
    # Check server health
    if not test_health():
        print("\n❌ Server is not running! Please start the backend.")
        return
    
    # Create test files
    print("\n📁 Creating test files...")
    pdf_path = create_test_pdf()
    doc_path = create_test_word()
    
    if not pdf_path:
        print("❌ Could not create test PDF")
        return
    if not doc_path:
        print("❌ Could not create test Word document")
        return
    
    print(f"   Created: {pdf_path}")
    print(f"   Created: {doc_path}")
    
    # Upload test files
    print("\n📤 Testing file uploads...")
    pdf_file_id = test_upload_pdf(pdf_path)
    word_file_id = test_upload_word(doc_path)
    
    if not pdf_file_id or not word_file_id:
        print("\n❌ Upload failed! Cannot continue tests.")
        return
    
    # PDF Tests
    print("\n📄 Testing PDF operations...")
    test_pdf_info(pdf_file_id)
    test_pdf_thumbnails(pdf_file_id)
    test_pdf_merge(pdf_file_id)
    test_pdf_split(pdf_file_id)
    test_pdf_rotate(pdf_file_id)
    test_pdf_compress(pdf_file_id)
    test_pdf_watermark(pdf_file_id)
    test_pdf_extract(pdf_file_id)
    test_pdf_delete(pdf_file_id)
    
    # PDF Protection tests
    print("\n🔒 Testing PDF protection...")
    protected_id = test_pdf_protect(pdf_file_id)
    if protected_id:
        test_pdf_unlock(protected_id)
    
    # Word Tests
    print("\n📝 Testing Word operations...")
    test_word_info(word_file_id)
    test_word_header(word_file_id)
    test_word_footer(word_file_id)
    test_word_replace(word_file_id)
    
    # Conversion Tests
    print("\n🔄 Testing conversions...")
    caps = test_convert_capabilities()
    test_word_to_pdf(word_file_id)
    test_pdf_to_images(pdf_file_id)
    
    # OCR Tests
    print("\n👁️ Testing OCR...")
    test_ocr_status()
    
    # Summary
    print("\n" + "=" * 60)
    print("TEST SUMMARY")
    print("=" * 60)
    print(f"✅ Passed: {len(results['passed'])}")
    print(f"❌ Failed: {len(results['failed'])}")
    
    if results['failed']:
        print("\n❌ Failed tests:")
        for test_name, error in results['failed']:
            print(f"   - {test_name}: {error}")
    
    print()


if __name__ == "__main__":
    main()
